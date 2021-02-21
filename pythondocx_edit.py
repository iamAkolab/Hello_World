import docx

mydocx = docx.Document('C:\\Users\\john\\Downloads\\demo.docx')
print(mydocx.paragraphs)

print(mydocx.paragraphs[0])
print(mydocx.paragraphs[0].text)

print(mydocx.paragraphs[1].text)

#a new run object is when there is a change in style
p = mydocx.paragraphs[1]
print(p.runs)

print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

#We can get whether it is bold or not bold italic
print(p.runs[0].bold)
print(p.runs[0].bold == None)


print(p.runs[3].italic)
print(p.runs[3].underline == True)
p.runs[3].text = 'italic and underline'

p.save('C:\\Users\\john\\Downloads\\demo2.docx')

print(p.style)
p.style = 'Title'
p.save('C:\\Users\\john\\Downloads\\demo3.docx')

#Create  a new doc that resides inside python
d = docx.Document()

d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph')

d.save('C:\\Users\\john\\Downloads\\demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run')

p.runs
p.runs[1].bold = True

d.save('C:\\Users\\john\\Downloads\\demo5.docx')
